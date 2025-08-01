import logging
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from openai import OpenAI
import os
import uuid
import json
import re
import io

# --- YEH SECTION UPDATE KIYA GAYA HAI ---
# Libraries for DOCX generation and PDF handling
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import PyPDF2
import pdfkit
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
# html2docx ki ab zaroorat nahi hai, isliye hata diya gaya hai
# --------------------------------------

# Local application imports from resume_ai_analyzer.py
from resume_ai_analyzer import (
    # New and Corrected Functions for the Stable Strategy
    extract_resume_sections_safely,
    generate_final_detailed_report,
    fix_resume_issue,
    calculate_new_score,
    get_field_suggestions,
    generate_smart_resume_from_keywords,
    generate_full_ai_resume_html, # <--- YE LINE ADD KAREIN

    # Existing Utility and Other Functions
    analyze_resume_with_openai,
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_resume,
    check_ats_compatibility,
    check_ats_compatibility_fast,
    check_ats_compatibility_deep,
    extract_keywords_from_jd,
    compare_resume_with_keywords,
    generate_keyword_suggestions,
    generate_resume_score_and_detailed_feedback,
    analyze_job_description,
    fix_resume_formatting,
    generate_resume_summary,
    generate_michelle_template_html
)
# NEW/FIXED: Correctly import html2docx at the top level
try:
    from html2docx import html2docx
except ImportError:
    logging.error("html2docx library is not installed or accessible. DOCX conversion will fail.")
    html2docx = None

logging_level = logging.INFO if os.environ.get("FLASK_ENV") != "development" else logging.DEBUG
logging.basicConfig(level=logging_level)
logger = logging.getLogger(__name__)
    
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError as e:
    logging.error(f"Failed to import python-docx: {str(e)}")
    raise

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
except ImportError as e:
    logging.error(f"Failed to import reportlab: {str(e)}")
    raise

try:
    import fitz  # PyMuPDF for PDF text extraction
    import PyPDF2  # For PDF validation (encryption check)
    import pdfkit  # For DOCX to PDF conversion
except ImportError as e:
    logging.error(f"Failed to import required dependencies: {str(e)}")
    raise
import io  # Required for BytesIO in DOCX conversion
import tempfile  # Required for temporary file handling

logging_level = logging.INFO if os.environ.get("FLASK_ENV") != "development" else logging.DEBUG
logging.basicConfig(level=logging_level)
logger = logging.getLogger(__name__)

logging.getLogger("pdfminer").setLevel(logging.ERROR)

logger.info("Starting Flask app initialization...")

app = Flask(__name__, static_url_path='/static')
CORS(app, resources={r"/*": {"origins": "https://resumefixerpro.com"}})

UPLOAD_FOLDER = '/tmp/Uploads'  # Use /tmp for Render compatibility
STATIC_FOLDER = '/tmp/static'
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(STATIC_FOLDER, exist_ok=True)
    logger.info(f"Directories created: {UPLOAD_FOLDER}, {STATIC_FOLDER}")
except Exception as e:
    logger.error(f"Failed to create directories: {str(e)}")
    raise RuntimeError(f"Failed to create directories: {str(e)}")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Removes unnecessary personal details from resume text
def remove_unnecessary_personal_info(text):
    lines = text.split('\n')
    filtered_lines = []
    flagged_terms = ['Marital Status', 'Date of Birth', 'Gender', 'Nationality', 'Religion']
    for line in lines:
        if not any(term.lower() in line.lower() for term in flagged_terms):
            filtered_lines.append(line)
    return '\n'.join(filtered_lines)


def cleanup_file(filepath):
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.debug(f"Cleaned up file: {filepath}")
    except Exception as e:
        logger.error(f"Error cleaning up file: {str(e)}")

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "message": "OK"}), 200

@app.route("/upload", methods=["POST"])
def upload_resume():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        # ✅ Extract text directly from uploaded file
        text = extract_text_from_resume(file)

        if not text.strip():
            return jsonify({"error": "Failed to extract resume text"}), 500

        # ✅ Smart section parser - this is already here, which is great!
        parsed_sections = extract_resume_sections(text)

        # ✅ Optional summary generation
        name = parsed_sections.get("name", "")
        role = parsed_sections.get("job_title", "")
        experience = parsed_sections.get("work_experience", "")
        skills = parsed_sections.get("skills", "")
        summary = generate_resume_summary(name, role, experience, skills)
        parsed_sections["summary"] = summary

        # ✅ ATS report and score (THIS IS THE CORRECTED LINE)
        ats_result = generate_ats_report(text, parsed_sections)
        ats_issues = ats_result["issues"]
        score = ats_result["score"]

        return jsonify({
            "resume_text": text,
            "parsedResumeContent": parsed_sections,
            "ats_report": ats_issues,
            "score": score
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@app.route('/main-upload', methods=['POST'])
def main_upload():
    try:
        logger.info("Received request for /main-upload")
        file = request.files.get('file')
        if not file or file.filename == '':
            logger.error("No file uploaded in request")
            return jsonify({"error": "No file uploaded"}), 400

        logger.info(f"File received: {file.filename}")
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in {'.pdf', '.docx'}:
            logger.error(f"Unsupported file format: {ext}")
            return jsonify({"error": f"Unsupported file format: {ext}. Please upload a PDF or DOCX."}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        logger.info(f"Saving file to {filepath}")
        file.save(filepath)

        # ✅ Extract original resume text (NO cleaning)
        resume_text = extract_text_from_pdf(filepath) if ext == ".pdf" else extract_text_from_docx(filepath)

        # ✅ Run ATS checks on raw text (personal info will be caught)
        ats_score, final_checks = perform_ats_checks(resume_text)

        # ✅ Extract section-wise smart content
        smart_content = generate_smart_content(resume_text)

        # ✅ Fallback: Manually extract contact if missing
        if not smart_content.get("contact"):
            import re
            email = re.search(r'[\w\.-]+@[\w\.-]+', resume_text)
            phone = re.search(r'(\+?\d[\d\s\-\(\)]{7,})', resume_text)
            location = re.search(r'\b(?:[A-Z][a-z]+\s?){1,3}(?:,?\s?(India|USA|UK|Canada|Australia))?', resume_text)
            smart_content["contact"] = {
                "email": email.group() if email else "",
                "phone": phone.group() if phone else "",
                "location": location.group() if location else ""
            }
            logger.info("Fallback contact used: %s", smart_content["contact"])

        # ✅ Format contact section properly
        contact_data = smart_content.get("contact", {})
        if isinstance(contact_data, dict):
            contact_lines = list(filter(None, [
                contact_data.get('email'),
                contact_data.get('phone'),
                contact_data.get('location'),
                contact_data.get('linkedin'),
                contact_data.get('github')
            ]))
            contact = "<ul>" + "".join(f"<li>{line}</li>" for line in contact_lines) + "</ul>" if contact_lines else ""
        else:
            contact = f"<p>{contact_data}</p>" if contact_data else ""

        # ✅ Final resume data with fallback values
        final_data = {
            "contact": contact,
            "skills": smart_content.get("skills", "") or "",
            "languages": smart_content.get("languages", "") or "",
            "certifications": smart_content.get("certifications", "") or "",
            "awards": smart_content.get("awards", "") or "",
            "summary": smart_content.get("summary", "") or "",
            "experience": smart_content.get("experience", "") or "",
            "education": smart_content.get("education", "") or "",
            "projects": smart_content.get("projects", "") or "",
            "publications": smart_content.get("publications", "") or "",
        }

        html = generateResumeTemplate(final_data)

        return jsonify({
            "success": True,
            "score": ats_score,
            "html": html,
            "checks": final_checks
        })

    except Exception as e:
        logger.error(f"Error in /main-upload: {str(e)}")
        return jsonify({"error": f"Failed to process resume: {str(e)}"}), 500

    finally:
        try:
            if 'filepath' in locals() and os.path.exists(filepath):
                cleanup_file(filepath)
        except Exception as cleanup_err:
            logger.warning(f"Cleanup failed: {cleanup_err}")

@app.route('/ats-check', methods=['POST'])
def check_ats():
    try:
        file = request.files.get('file')  # Only look for the key "file"
        if not file or file.filename == '':
            return jsonify({"error": "No file uploaded"}), 400

        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in {'.pdf', '.docx'}:
            return jsonify({"error": f"Unsupported file format: {ext}. Please upload a PDF or DOCX."}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        text = extract_text_from_pdf(filepath) if ext == ".pdf" else extract_text_from_docx(filepath)
        cleaned_text = remove_unnecessary_personal_info(text)

        # Enhanced ATS prompt with personal info warning
        prompt = f"""
You are an ATS expert. Review the following resume and identify up to 5 ATS-related issues.

⚠️ Also flag any unnecessary personal information such as:
- Marital Status
- Date of Birth
- Gender
- Nationality
- Religion

These are not required in a professional resume and should be removed for better ATS compatibility.

Resume:
{cleaned_text[:6000]}

Return your output as a list like this:
["✅ Passed: ...", "❌ Issue: ..."]
Only include meaningful points.
"""

        from openai import OpenAI
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        ai_resp = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )

        feedback = ai_resp.choices[0].message.content.strip().splitlines()

        # Normalize feedback
        formatted_feedback = []
        for line in feedback:
            line = line.strip()
            if not line:
                continue
            if line.lower().startswith("issue:"):
                formatted_feedback.append("❌ " + line[len("issue:"):].strip())
            elif line.lower().startswith("passed:"):
                formatted_feedback.append("✅ " + line[len("passed:"):].strip())
            elif line.startswith("❌") or line.startswith("✅"):
                formatted_feedback.append(line)
            else:
                formatted_feedback.append("❌ " + line)

        # Score calculation
        score = 100 - (len([line for line in formatted_feedback if line.startswith("❌")]) * 20)
        score = max(0, min(score, 100))

        return jsonify({"issues": formatted_feedback, "score": score})

    except Exception as e:
        logger.error(f"Error in /ats-check: {str(e)}")
        return jsonify({"error": f"Failed to check ATS compatibility: {str(e)}"}), 500

    finally:
        try:
            if 'filepath' in locals() and os.path.exists(filepath):
                cleanup_file(filepath)
        except Exception as cleanup_err:
            logger.warning(f"Cleanup failed: {cleanup_err}")

@app.route('/analyze', methods=['POST'])
def analyze_resume():
    try:
        data = request.get_json()
        resume_text = data.get('resume_text')
        if not resume_text or not isinstance(resume_text, str) or not resume_text.strip():
            return jsonify({"error": "Invalid or empty resume text"}), 400
        result = analyze_resume_with_openai(resume_text, atsfix=False)
        if "error" in result:
            return jsonify({"error": "Unable to generate suggestions. Please check if the API key is set."}), 500
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error in /analyze: {str(e)}")
        return jsonify({"error": f"Failed to analyze resume: {str(e)}"}), 500

@app.route('/fix-suggestion', methods=['POST'])
def fix_suggestion():
    try:
        if 'payload' not in request.form:
            return jsonify({"success": False, "error": "Missing payload in request form"}), 400
            
        data = json.loads(request.form.get('payload'))
        suggestion = data.get("suggestion")
        full_text = data.get("full_text")
        current_score = data.get("current_score") # <<< हमें JS से वर्तमान स्कोर चाहिए

        if not all([suggestion, full_text, isinstance(current_score, int)]):
            return jsonify({"success": False, "error": "Missing suggestion, text, or current_score."}), 400

        # STEP 1: Get the targeted fix from the AI.
        fix_result = generate_targeted_fix(suggestion, full_text) # <<< बदला हुआ
        
        if 'error' in fix_result:
            return jsonify({"success": False, "error": fix_result["error"]}), 500

        # STEP 2: Calculate the new score predictably.
        new_score = calculate_new_score(current_score, suggestion) # <<< बदला हुआ
        
        # Combine results into the final response
        final_response = {
            "section": fix_result["section"],
            "fixedContent": fix_result["fixedContent"],
            "newScore": new_score
        }

        # NOTE: We no longer return 'updatedAnalysis'. The frontend will handle the UI state.
        return jsonify({"success": True, "data": final_response})

    except Exception as e:
        import traceback
        logger.error(f"Error in /fix-suggestion: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "An unexpected server error occurred."}), 500

@app.route('/preview-resume', methods=['POST'])
def preview_resume():
    try:
        data = request.get_json()
        sections = data.get('sections')
        if not sections:
            logger.error("No sections provided in /preview-resume request")
            return jsonify({"error": "No sections provided"}), 400

        # Generate HTML
        html_content = """
        <div style='width: 90%; margin: 0 auto; font-family: Arial, sans-serif;'>
        """
        for section, content in sections.items():
            section_title = section.replace('_', ' ').title()
            html_content += """
            <div class='section' style='margin-bottom: 1.2rem;'>
              <h3 style='font-size: 0.95rem; line-height: 1.3; color: #222; margin-bottom: 4px; border-bottom: 1px solid #ccc;'>{}</h3>
              <div style='line-height: 1.5;'>{}</div>
            </div>
            """.format(section_title.upper(), content.replace('\n', '<br>'))
        html_content += "</div>"

        # Generate plain text
        resume_text_lines = []
        for section, content in sections.items():
            section_title = section.replace('_', ' ').title()
            resume_text_lines.append(section_title.upper())
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    if not line.startswith('- '):
                        line = '- ' + line
                    resume_text_lines.append(line)
            resume_text_lines.append('')

        # Clean up excessive blank lines
        cleaned_resume = []
        last_was_empty = False
        for line in resume_text_lines:
            if line == '' and last_was_empty:
                continue
            cleaned_resume.append(line)
            last_was_empty = (line == '')
        resume_text = '\n'.join(cleaned_resume).strip()

        return jsonify({"preview_text": resume_text, "preview_html": html_content})

    except Exception as e:
        logger.error(f"Error in /preview-resume: {str(e)}")
        return jsonify({"error": f"Failed to generate preview: {str(e)}"}), 500

@app.route('/final-resume', methods=['POST'])
def final_resume_download():
    try:
        data = request.get_json()
        html_content = data.get("html")
        file_format = data.get("format", "pdf")

        if not html_content:
            logger.error("No HTML content provided in /final-resume request")
            return jsonify({"error": "Invalid HTML content provided"}), 400

        wrapped_html = """
        <div style='width: 90%; margin: 0 auto; font-family: Arial, sans-serif; padding: 20px;'>
            {}
        </div>
        """.format(html_content)

        if file_format in ["pdf", "docx"]:
            if file_format == "pdf":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_file:
                    pdfkit.from_string(wrapped_html, pdf_file.name)
                    logger.info(f"Generated PDF at: {pdf_file.name}")
                    return send_file(
                        pdf_file.name,
                        as_attachment=True,
                        download_name="Final_Resume.pdf",
                        mimetype='application/pdf'
                    )

            else:
                try:
                    from html2docx import html2docx
                except ImportError as e:
                    logger.error(f"html2docx not installed: {str(e)}")
                    return jsonify({"error": "Server error: html2docx not installed"}), 500

                docx_bytes = html2docx(wrapped_html)
                logger.info("Generated DOCX in memory")
                return send_file(
                    io.BytesIO(docx_bytes),
                    as_attachment=True,
                    download_name="Final_Resume.docx",
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

        else:
            logger.error(f"Invalid format requested: {file_format}")
            return jsonify({"error": "Invalid format. Use 'pdf' or 'docx'."}), 400

    except Exception as e:
        logger.error(f"Error in /final-resume: {str(e)}")
        return jsonify({"error": f"Failed to generate final resume: {str(e)}"}), 500

@app.route('/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    file = request.files.get('file') or request.files.get('resume')
    job_title = request.form.get('job_title')
    company_name = request.form.get('company_name')

    if not file or not job_title or not company_name:
        return jsonify({"error": "File, job title, and company name are required"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    try:
        file.save(filepath)
        resume_text = extract_text_from_resume(file)
        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from resume"}), 400

        prompt = """
You are a professional cover letter writer. Write a concise cover letter (300-400 words) for the following details:
Job Title: {}
Company Name: {}
Resume: {}
Include a greeting, an introduction, a body highlighting relevant skills and experiences, and a closing statement.
        """.format(job_title, company_name, resume_text[:6000])

        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            cover_letter = response.choices[0].message.content.strip()
            return jsonify({"cover_letter": cover_letter})
        except Exception as e:
            logger.error(f"Error in OpenAI API call for /generate-cover-letter: {str(e)}")
            return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500
    except Exception as e:
        logger.error(f"Error in /generate-cover-letter: {str(e)}")
        return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500
    finally:
        cleanup_file(filepath)
        
@app.route('/generate-cover-letter-from-data', methods=['POST'])
def generate_cover_letter_from_data():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data provided"}), 400

        extracted_data = data.get('extracted_data')
        job_title = data.get('job_title')
        company_name = data.get('company_name')

        if not extracted_data or not job_title or not company_name:
            return jsonify({"error": "extracted_data, job_title, and company_name are required"}), 400

        # extracted_data (JSON) se ek plain text resume banayein
        resume_text_parts = []
        if extracted_data.get('name'):
            resume_text_parts.append(f"Name: {extracted_data.get('name')}")
        if extracted_data.get('summary'):
            resume_text_parts.append(f"Summary: {extracted_data.get('summary')}")
        if extracted_data.get('work_experience'):
            resume_text_parts.append("\nWork Experience:")
            for exp in extracted_data.get('work_experience', []):
                resume_text_parts.append(f"- {exp.get('title')} at {exp.get('company')}")
        if extracted_data.get('skills'):
            skills_str = ', '.join(extracted_data.get('skills', []))
            resume_text_parts.append(f"\nSkills: {skills_str}")
        
        resume_text = "\n".join(resume_text_parts)

        if not resume_text.strip():
            return jsonify({"error": "Could not construct resume text from provided data"}), 400

        prompt = f"""
You are a professional cover letter writer. Write a concise cover letter (300-400 words) for the following details:
Job Title: {job_title}
Company Name: {company_name}
Resume Highlights: {resume_text[:6000]}
Include a greeting, an introduction, a body highlighting relevant skills and experiences, and a closing statement.
        """

        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            cover_letter = response.choices[0].message.content.strip()
            return jsonify({"cover_letter": cover_letter})
        except Exception as e:
            logger.error(f"Error in OpenAI API call for /generate-cover-letter-from-data: {str(e)}")
            return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"Error in /generate-cover-letter-from-data: {str(e)}")
        return jsonify({"error": "An unexpected server error occurred"}), 500
    
@app.route('/download-cover-letter', methods=['POST'])
def download_cover_letter():
    data = request.get_json()
    cover_letter = data.get('cover_letter')
    if not cover_letter:
        return jsonify({"error": "No cover letter provided"}), 400

    output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"cover_letter_{uuid.uuid4()}.docx")
    try:
        doc = Document()
        doc.add_heading("Cover Letter", level=1)
        for line in cover_letter.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.save(output_path)
        return send_file(
            output_path,
            as_attachment=True,
            download_name="Cover_Letter.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Error in /download-cover-letter: {str(e)}")
        return jsonify({"error": f"Failed to download cover letter: {str(e)}"}), 500
    finally:
        cleanup_file(output_path)

@app.route('/resume-score', methods=['POST'])
def resume_score():
    resume_text = ""
    filepath = None  # For cleanup

    try:
        if request.is_json:
            data = request.get_json()
            resume_text = data.get("resume_text", "")
        else:
            file = request.files.get('file')
            if not file:
                return jsonify({"error": "No file uploaded"}), 400
            
            # File ko save karein aur text extract karein
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            ext = os.path.splitext(filepath)[1].lower()
            if ext == ".pdf":
                resume_text = extract_text_from_pdf(filepath)
                if resume_text == "NO_TEXT_EXTRACTED_IMAGE_BASED":
                    return jsonify({"error": "Please upload a scannable PDF. Image-based PDFs are not supported."}), 400
            elif ext == ".docx":
                resume_text = extract_text_from_docx(filepath)
            else:
                return jsonify({"error": "Unsupported file format"}), 400

        if not resume_text.strip():
            return jsonify({"error": "No extractable text found in resume"}), 400

        # Naye function ko call karein jo score aur detailed feedback dega
        # generate_resume_score_and_detailed_feedback function ko resume_ai_analyzer.py se import karna zaroori hai
        report = generate_resume_score_and_detailed_feedback(resume_text)

        if "error" in report:
            logger.error(f"Error generating detailed report for score checker: {report['error']}")
            return jsonify({"error": report["error"]}), 500

        # Frontend ko detailed score aur issues bhej dein
        return jsonify({
            "score": report["score"],
            "feedback_details": report["feedback_details"] # Detailed issues aur suggestions
        })

    except Exception as e:
        import traceback
        logger.error(f"Error in /resume-score: {traceback.format_exc()}")
        # Fallback score aur generic error message agar kuch bhi galat hota hai
        return jsonify({"score": 70, "feedback_details": [{"status": "error", "comment": f"An unexpected error occurred during analysis: {str(e)}"}]})
    finally:
        if filepath:
            cleanup_file(filepath) # Temporary file cleanup

@app.route('/optimize-keywords', methods=['POST'])
def optimize_keywords():
    resume_file = None
    filepath = None # For cleanup

    try:
        resume_file = request.files.get('resume')
        job_description_text = request.form.get('job_description', '')

        if not resume_file or resume_file.filename == '':
            logger.error("No resume file uploaded for /optimize-keywords")
            return jsonify({"error": "Missing resume file."}), 400
        
        if not job_description_text.strip():
            logger.error("No job description provided for /optimize-keywords")
            return jsonify({"error": "Missing job description."}), 400

        # Resume file ko temp location par save karein
        filename = secure_filename(resume_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        resume_file.save(filepath)
        logger.info(f"Resume file saved to: {filepath}")

        # Step 1: Resume se text extract karein
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            resume_text = extract_text_from_pdf(filepath)
            # Image-based PDF detection
            if resume_text == "NO_TEXT_EXTRACTED_IMAGE_BASED":
                logger.warning(f"Image-based PDF detected for {filename}.")
                return jsonify({"error": "Please upload a scannable PDF. We currently do not support image-based PDFs."}), 400
        elif ext == ".docx":
            resume_text = extract_text_from_docx(filepath)
        else:
            logger.error(f"Unsupported resume file format: {ext}")
            return jsonify({"error": f"Unsupported resume file format: {ext}. Please upload a PDF or DOCX."}), 400

        if not resume_text.strip():
            logger.error(f"No extractable text found in resume: {filename}")
            return jsonify({"error": "Could not extract text from your resume. It might be empty or image-based."}), 400

        # Step 2: Job Description se keywords extract karein
        jd_keywords_string = extract_keywords_from_jd(job_description_text)
        if not jd_keywords_string.strip():
            logger.error("No keywords extracted from job description by AI.")
            return jsonify({"error": "Could not extract keywords from the job description. Please ensure it contains relevant text."}), 400
        
        # Step 3: Resume aur JD keywords ko compare karein
        # compare_resume_with_keywords ab 'present_keywords' return karega.
        comparison_results = compare_resume_with_keywords(resume_text, jd_keywords_string)
        
        present_keywords = comparison_results.get("present_keywords", [])
        missing_keywords = comparison_results.get("missing_keywords", [])
        
        # Step 4: Missing keywords ke liye AI suggestions generate karein
        # Ye naya function hai jo resume_ai_analyzer.py mein add kiya jayega.
        suggested_keywords = generate_keyword_suggestions(job_description_text, missing_keywords)
        
        # Final response frontend ko bhej dein
        return jsonify({
            "match_score": comparison_results.get("match_score", 0),
            "present_keywords": present_keywords,
            "missing_keywords": missing_keywords,
            "suggested_keywords": suggested_keywords
        })

    except Exception as e:
        import traceback
        logger.error(f"Error in /optimize-keywords: {traceback.format_exc()}")
        return jsonify({"error": f"An unexpected server error occurred: {str(e)}"}), 500
    finally:
        if filepath and os.path.exists(filepath):
            os.remove(filepath) # Temporary file cleanup
            logger.debug(f"Cleaned up temporary file: {filepath}")

@app.route('/analyze-jd', methods=['POST'])
def analyze_job_description_api():
    try:
        data = request.get_json()
        jd_text = data.get('job_description')

        if not jd_text or not isinstance(jd_text, str) or not jd_text.strip():
            logger.error("Invalid or empty job description text provided for /analyze-jd")
            return jsonify({"error": "Invalid or empty job description"}), 400
        
        # Call the updated analyze_job_description function from resume_ai_analyzer.py
        # This function will now directly return a JSON object
        analysis_result = analyze_job_description(jd_text) 

        if "error" in analysis_result:
            logger.error(f"Error from analyze_job_description: {analysis_result['error']}")
            return jsonify({"error": analysis_result["error"]}), 500

        # Return the JSON response directly as analyze_job_description is now in the correct format
        return jsonify(analysis_result)

    except Exception as e:
        logger.error(f"Error in /analyze-jd API: {str(e)}")
        return jsonify({"error": f"Failed to analyze job description: {str(e)}"}), 500

# app.py में, analyze_job_description_api फंक्शन के बाद इस नए राउट को जोड़ें YE jd ANALYSE KA HI HAI PDF KE LIYE

@app.route('/analyze-jd-file', methods=['POST'])
def analyze_jd_from_file_api():
    try:
        file = request.files.get('file')
        if not file or file.filename == '':
            logger.error("No file uploaded for /analyze-jd-file")
            return jsonify({"error": "No file uploaded"}), 400

        # Save file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text from the file
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            jd_text = extract_text_from_pdf(filepath)
        elif ext == ".docx":
            jd_text = extract_text_from_docx(filepath)
        else:
            cleanup_file(filepath)
            return jsonify({"error": "Unsupported file format. Please upload PDF or DOCX."}), 400

        if not jd_text.strip():
            cleanup_file(filepath)
            return jsonify({"error": "Could not extract text from the uploaded file. File might be empty or image-based."}), 400

        # Call the analyze_job_description function (जो अब JSON रिटर्न करता है)
        analysis_result = analyze_job_description(jd_text)

        if "error" in analysis_result:
            logger.error(f"Error from analyze_job_description (file upload): {analysis_result['error']}")
            return jsonify({"error": analysis_result["error"]}), 500

        return jsonify(analysis_result)

    except Exception as e:
        logger.error(f"Error in /analyze-jd-file API: {str(e)}")
        return jsonify({"error": f"Failed to analyze job description from file: {str(e)}"}), 500
    finally:
        # Clean up the temporary file
        if 'filepath' in locals() and os.path.exists(filepath):
            cleanup_file(filepath)

# Yeh function app.py me replace karein
@app.route("/generate-ai-resume", methods=["POST"])
def generate_ai_resume():
    try:
        data = request.get_json()
        contact_fields = ["name", "email", "phone", "location", "linkedin", "jobTitle"]
        user_info = {key: data.get(key, "") for key in contact_fields}
        section_data = {key: value for key, value in data.items() if key not in contact_fields}

        # Generate smart content from keywords
        smart_content = generate_smart_resume_from_keywords(section_data)

        # Generate HTML with updated function
        html = generate_full_ai_resume_html(user_info, smart_content)

        return jsonify({"success": True, "html": html})
    except Exception as e:
        import traceback
        traceback.print_exc()  # Debugging ke liye traceback print karen
        return jsonify({"error": f"❌ Exception in generate-ai-resume: {type(e).__name__} - {str(e)}"}), 500

# app.py file mein, convert_format route ke andar ke changes:

@app.route('/convert-format', methods=['POST'])
def convert_format():
    try:
        file = request.files.get('file')
        target_format = request.form.get('target_format')

        if not file or not target_format:
            logger.error("Missing file or target format in request")
            return jsonify({"error": "Missing file or target format specified"}), 400

        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in ['.pdf', '.docx']:
            logger.error(f"Unsupported file format: {ext}")
            return jsonify({"error": f"Invalid file format: {ext}. Please upload a PDF or DOCX file."}), 400

        filename = secure_filename(file.filename)
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"converted_{uuid.uuid4()}.{target_format}")
        html_temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{uuid.uuid4()}.html")

        file.save(upload_path)
        logger.debug(f"Saved uploaded file to {upload_path}")

        if not os.path.exists(upload_path):
            logger.error(f"File not found after saving: {upload_path}")
            return jsonify({"error": "File could not be saved properly"}), 500

        os.chmod(upload_path, 0o644)
        file_size = os.path.getsize(upload_path) / 1024
        if file_size == 0:
            logger.error(f"Empty file uploaded: {filename}")
            return jsonify({"error": "Uploaded file is empty"}), 400

        logger.debug(f"File size: {file_size:.2f} KB")

        if ext == '.pdf':
            try:
                with open(upload_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    if pdf_reader.is_encrypted:
                        logger.error("Uploaded PDF is encrypted")
                        return jsonify({"error": "PDF is encrypted and cannot be processed."}), 400
            except Exception as e:
                logger.error(f"Invalid or corrupted file: {str(e)}")
                return jsonify({"error": f"Invalid or corrupted file: {str(e)}"}), 400

        if ext == '.docx':
            try:
                doc = Document(upload_path)
            except Exception as e:
                logger.error(f"Invalid or corrupted DOCX file: {str(e)}")
                return jsonify({"error": f"Invalid or corrupted DOCX file: {str(e)}"}), 400

        # --- Common Text Extraction Logic ---
        # extract_text_from_pdf function ko resume_ai_analyzer.py mein update kiya gaya hai.
        extracted_text = ""
        if ext == '.pdf':
            extracted_text = extract_text_from_pdf(upload_path)
            if extracted_text == "NO_TEXT_EXTRACTED_IMAGE_BASED":
                return jsonify({"error": "Please upload a scannable PDF. We currently do not support image-based PDFs."}), 400
        elif ext == '.docx':
            # DOCX se text extract karte waqt, paragraphs ko preserve karein
            doc = Document(upload_path) # Re-open doc for consistent processing
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
            
        if not extracted_text.strip():
            logger.warning("No text could be extracted from the file")
            return jsonify({"error": "No text could be extracted. The file may be empty or contain only images."}), 400

        # --- Conversion Logic based on target_format ---
        if target_format == 'text':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            
            return send_file(
                output_path,
                as_attachment=True,
                download_name="extracted-text.txt",
                mimetype='text/plain'
            )

        elif ext == '.docx' and target_format == 'pdf':
            try:
                doc = Document(upload_path) # Load DOCX for style analysis
                html_content_parts = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    # DOCX styles ko HTML headings aur lists mein map karein
                    style_name = para.style.name.lower()
                    if "heading 1" in style_name:
                        html_content_parts.append(f"<h1>{text}</h1>")
                    elif "heading 2" in style_name:
                        html_content_parts.append(f"<h2>{text}</h2>")
                    elif "heading 3" in style_name:
                        html_content_parts.append(f"<h3>{text}</h3>")
                    elif "list" in style_name or text.startswith(('-', '*', '•')) or (len(text) > 2 and text[0].isdigit() and text[1] == '.'):
                        # Bullet/Numbered list items
                        html_content_parts.append(f"<li>{text.lstrip('-*•0-9. ').strip()}</li>")
                    else:
                        html_content_parts.append(f"<p>{text}</p>")

                if not html_content_parts:
                    logger.error("DOCX file has no readable content for PDF conversion")
                    return jsonify({"error": "DOCX file is empty or contains no readable text."}), 400

                # List items ko <ul> tags mein wrap karein
                final_html_content = ""
                in_list = False
                for part in html_content_parts:
                    if part.startswith("<li>") and not in_list:
                        final_html_content += "<ul>"
                        in_list = True
                    elif not part.startswith("<li>") and in_list and not part.startswith("<li>"): # List end ho gayi
                        final_html_content += "</ul>"
                        in_list = False
                    final_html_content += part
                if in_list: # Agar list end nahi hui thi
                    final_html_content += "</ul>"

                # Basic styling for PDF
                html_for_pdf = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: 'sans-serif'; margin: 20px; line-height: 1.6; color: #333; }}
                        h1 {{ font-size: 24px; margin-top: 20px; margin-bottom: 10px; }}
                        h2 {{ font-size: 20px; margin-top: 18px; margin-bottom: 8px; }}
                        h3 {{ font-size: 18px; margin-top: 15px; margin-bottom: 7px; }}
                        p {{ margin-bottom: 10px; }}
                        ul {{ margin-left: 20px; margin-bottom: 10px; list-style-type: disc; }}
                        li {{ margin-bottom: 5px; }}
                    </style>
                </head>
                <body>
                    {final_html_content}
                </body>
                </html>
                """
                
                with open(html_temp_path, 'w', encoding='utf-8') as f:
                    f.write(html_for_pdf)
                
                pdfkit.from_file(html_temp_path, output_path)
                
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name="converted.pdf",
                    mimetype='application/pdf'
                )

            except Exception as e:
                logger.error(f"DOCX to PDF conversion failed: {str(e)}")
                return jsonify({"error": f"Failed to convert DOCX to PDF: {str(e)}"}), 500

        elif ext == '.pdf' and target_format == 'docx':
            try:
                # PDF se text extract karein (extracted_text variable upar se use hoga)
                word_doc = Document()
                
                # Heuristics for PDF text to DOCX formatting
                lines = extracted_text.split('\n')
                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue # Skip empty lines

                    # Simple Heading Detection (Heuristic-based)
                    # Check for all caps, short lines, or common heading patterns
                    is_heading = False
                    if len(line) < 50 and line.isupper() and len(line.split()) < 5:
                        is_heading = True
                    elif re.match(r"^(education|experience|skills|summary|projects|certifications|languages|awards|volunteer|publications)\b", line.lower()):
                        is_heading = True
                    
                    # Bullet Point Detection
                    is_bullet = False
                    if line.startswith(('-', '*', '•')) or (len(line) > 2 and line[0].isdigit() and line[1] == '.' and ' ' in line):
                        is_bullet = True
                    
                    if is_heading:
                        # Attempt to assign heading levels based on length/position (simple heuristic)
                        if len(line) < 20 and line.isupper(): # Main sections
                            word_doc.add_heading(line, level=1)
                        else: # Sub-sections
                            word_doc.add_heading(line, level=2)
                    elif is_bullet:
                        word_doc.add_paragraph(line.lstrip('-*•0-9. ').strip(), style='List Bullet')
                    else:
                        word_doc.add_paragraph(line) # Normal paragraph

                if not word_doc.paragraphs:
                    logger.warning("No content added to DOCX from PDF extraction.")
                    return jsonify({"error": "No readable content found in PDF for DOCX conversion."}), 400

                word_doc.save(output_path)
                logger.info(f"Converted PDF to DOCX: {output_path}")
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name="converted.docx",
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            except Exception as e:
                logger.error(f"Failed to convert PDF to DOCX: {str(e)}")
                return jsonify({"error": f"Failed to convert PDF to DOCX: {str(e)}"}), 500

        else:
            logger.error(f"Unsupported conversion request: {ext} to {target_format}")
            return jsonify({"error": "Unsupported conversion request. Only PDF to DOCX, DOCX to PDF, or text extraction are supported."}), 400

    except Exception as e:
        logger.error(f"Error in /convert-format: {str(e)}")
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 500
    finally:
        cleanup_file(upload_path)
        cleanup_file(output_path)
        cleanup_file(html_temp_path) # html_temp_path ko bhi cleanup karein

@app.route('/fix-formatting', methods=['POST'])
def fix_formatting():
    file = request.files.get('file') or request.files.get('resume')
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        file.save(filepath)
        result = fix_resume_formatting(filepath)
        logger.info(f"Formatted resume at: {filepath}")
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error in /fix-formatting: {str(e)}")
        return jsonify({"error": f"Failed to process resume formatting: {str(e)}"}), 500
    finally:
        cleanup_file(filepath)

@app.route('/generate-resume-summary', methods=['POST'])
def generate_resume_summary_api():
    try:
        data = request.get_json()
        name = data.get('name', '')
        role = data.get('role', '')
        experience = data.get('experience', '')
        skills = data.get('skills', '')

        if not name or not role or not experience or not skills:
            return jsonify({"error": "Missing required fields"}), 400

        summary = generate_resume_summary(name, role, experience, skills)
        return jsonify({"summary": summary})

    except Exception as e:
        logger.error(f"Error in /generate-resume-summary: {str(e)}")
        return jsonify({"error": f"Failed to generate resume summary: {str(e)}"}), 500

@app.route('/send-feedback', methods=['POST'])
def send_feedback():
    try:
        data = request.form
        name = data.get('name', 'Unknown')
        email = data.get('email', '')
        message = data.get('message', '')
        file = request.files.get('screenshot')

        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.base import MIMEBase
        from email import encoders
        import smtplib

        sender_email = "help@resumefixerpro.com"
        receiver_email = "help@resumefixerpro.com"
        smtp_server = "smtp.hostinger.com"
        smtp_port = 587
        smtp_password = os.environ.get("SMTP_PASSWORD")

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = 'Feedback Submission from ResumeFixerPro'

        body = """
Name: {}
Email: {}
Message:
{}
""".format(name, email, message)
        msg.attach(MIMEText(body, 'plain'))

        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            with open(filepath, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())

            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={filename}')
            msg.attach(part)

            cleanup_file(filepath)

        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, smtp_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()

        return jsonify({"success": True, "message": "Feedback sent successfully!"})

    except Exception as e:
        logger.error(f"Error sending feedback: {str(e)}")
        return jsonify({"error": f"Failed to send feedback: {str(e)}"}), 500

@app.route('/send-message', methods=['POST'])
def send_message():
    try:
        data = request.get_json()
        name = data.get('name', 'Unknown')
        email = data.get('email', '')
        message = data.get('message', '')

        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        import smtplib

        sender_email = "help@resumefixerpro.com"
        receiver_email = "help@resumefixerpro.com"
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        smtp_password = os.environ.get("SMTP_PASSWORD")

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = 'Contact Message from ResumeFixerPro'

        body = """
New Contact Message:
Name: {}
Email: {}
Message:
{}
""".format(name, email, message)
        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, smtp_password)
        server.send_message(msg)
        server.quit()

        return jsonify({"success": True, "message": "Message sent successfully!"})

    except Exception as e:
        logger.error(f"Error sending message: {str(e)}")
        return jsonify({"error": f"Failed to send message: {str(e)}"}), 500
    
@app.route('/ask-ai', methods=['POST', 'OPTIONS'])
def ask_ai_handler():
    # CORS preflight request ko handle karna
    if request.method == 'OPTIONS':
        return '', 200
    
    # Asli POST request ko handle karna
    if request.method == 'POST':
        try:
            # --- OPENAI CLIENT FUNCTION KE ANDAR SET KIYA GAYA HAI ---
            # Yeh 'client is not defined' error ko theek karega
            client = OpenAI(
                api_key=os.environ.get("OPENAI_API_KEY"),
            )
            # ---------------------------------------------------------

            data = request.get_json()
            if not data or 'question' not in data:
                return jsonify({"success": False, "error": "Question is missing in the payload"}), 400

            question = data['question']

            # --- SMART SYSTEM PROMPT ---
            system_prompt = """
            You are 'ProBot', the friendly and helpful AI assistant for ResumeFixerPro.com.
            Your primary goal is to assist users with questions about the website's tools and provide helpful, resume-related advice.
            Your personality is encouraging, professional, and supportive. Always keep your answers concise.

            Here is the key information about ResumeFixerPro.com you must know:

            **About the Website:**
            - Founder: Imran Ali
            - Country of Origin: India
            - Purpose: The main goal is to provide powerful, high-quality resume tools for free to help everyone, from students to professionals, in their career journey.

            **Our Tools (9 in total):**
            1.  AI Resume Generator
            2.  Cover Letter Generator
            3.  Resume Score Checker
            4.  ATS Compatibility Checker
            5.  Format Converter
            6.  Formatting Fixer
            7.  Keyword Optimizer
            8.  Job Description Analyzer
            9.  Summary Generator

            **Contact Information:**
            - For any detailed help, users can visit the contact page: https://resumefixerpro.com/contact-us/
            
            When a user asks a question, use this information to provide an accurate response.
            """

            # OpenAI se jawab paane ke liye
            completion = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": question}
                ]
            )
            
            ai_answer = completion.choices[0].message.content
            
            # Jawab ko JSON format mein wapas bhejna
            return jsonify({"success": True, "answer": ai_answer})

        except Exception as e:
            logging.error(f"Error in /ask-ai: {str(e)}")
            return jsonify({"success": False, "error": "An error occurred while processing your question."}), 500
        
@app.route('/extract-sections', methods=['POST'])
def extract_sections():
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text.strip():
            logger.warning("No resume text provided in /extract-sections request")
            return jsonify({"error": "No resume text provided"}), 400

        # ✅ Smart section extraction
        sections = extract_resume_sections(text)

        # ✅ Inject fallback empty keys if missing (optional)
        required_keys = ["name", "job_title", "contact", "summary", "education", "work_experience", "projects", "skills", "certifications", "languages", "linkedin", "achievements"]
        for key in required_keys:
            if key not in sections:
                sections[key] = ""

        # ✅ Check if anything found at all
        if not any(sections.values()):
            logger.warning("No sections could be extracted from the resume text")
            return jsonify({"error": "Failed to extract sections: Resume format may be unsupported or unstructured."}), 400

        logger.info(f"Successfully extracted sections: {list(sections.keys())}")
        return jsonify(sections)

    except Exception as e:
        logger.error(f"Error extracting sections: {str(e)} | Resume text: {text[:500]}")
        return jsonify({"error": f"Failed to extract sections: {str(e)}"}), 500

# =====================================================================
# FINAL API ENDPOINT FOR THE NEW WORDPRESS FRONTEND (CORRECTED INDENTATION)
# =====================================================================

@app.route('/api/v1/analyze-resume', methods=['POST'])
def analyze_resume_for_frontend():
    try:
        if 'resume_file' not in request.files:
            return jsonify({"success": False, "error": "No 'resume_file' part in the request."}), 400
        
        file = request.files['resume_file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected."}), 400

        text = extract_text_from_resume(file)
        if not text:
            return jsonify({"success": False, "error": "Could not extract text from the resume."}), 500

        # Step 1: Extract sections
        extracted_sections = extract_resume_sections_safely(text)
        if not extracted_sections or extracted_sections.get("error"):
            error_message = extracted_sections.get("error", "Failed to parse resume sections.")
            logger.error(f"Error from extract_resume_sections_safely: {error_message}")
            return jsonify({"success": False, "error": error_message}), 500

        # Step 2: Generate detailed ATS report
        ats_result = generate_final_detailed_report(extracted_sections)
        if not ats_result or ats_result.get("error"):
            error_message = ats_result.get("error", "Failed to generate ATS report.")
            logger.error(f"Error from generate_final_detailed_report: {error_message}")
            return jsonify({"success": False, "error": error_message}), 500
        
        # Step 3: Get field-aware suggestions
        field_info = get_field_suggestions(extracted_sections, text)
        if not field_info or field_info.get("error"):
            logger.warning("Could not get field suggestions, proceeding with default.")
            field_info = {"field": "General", "suggestions": []}

        # Step 4: Prepare the final data structure
        formatted_data = {
            "analysis": ats_result,
            "extracted_data": extracted_sections,
            "field_info": field_info
        }

        # Step 5: Calculate the score based on the detailed report
        fail_count = sum(1 for check in formatted_data['analysis'].values() if isinstance(check, dict) and check.get('status') in ['fail', 'improve'])
        formatted_data['score'] = max(40, 100 - (fail_count * 10))
        
        return jsonify({"success": True, "data": formatted_data})

    except Exception as e:
        import traceback
        logger.error(f"Error in /api/v1/analyze-resume: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "An unexpected server error occurred."}), 500
from docx import Document
import html2text

@app.route('/api/v1/fix-issue-v2', methods=['POST'])
def handle_fix_issue_v2():
    try:
        if 'payload' not in request.form:
            logger.error("Payload missing in request for /api/v1/fix-issue-v2")
            return jsonify({"success": False, "error": "Missing payload"}), 400

        payload = json.loads(request.form.get('payload'))
        issue_text = payload.get('issue_text')
        extracted_data = payload.get('extracted_data')
        current_score = payload.get('current_score', 60) # Get current score, with a default

        if not issue_text or not extracted_data:
            logger.error("Missing 'issue_text' or 'extracted_data' in payload.")
            return jsonify({"success": False, "error": "Missing required data in payload"}), 400

        # Naye, reliable function ko call karein
        logger.info(f"Calling fix_resume_issue for: {issue_text[:80]}")
        # Make sure you have imported fix_resume_issue at the top of app.py
        fix_result = fix_resume_issue(issue_text, extracted_data) 

        if 'error' in fix_result:
            logger.error(f"fix_resume_issue function failed: {fix_result['error']}")
            return jsonify({"success": False, "error": fix_result['error']}), 500
        
        # Ek simple score badhane ka logic
        new_score = min(100, int(current_score) + 8)

        # Frontend ke liye response tayyar karein
        response_data = {
            "fix": fix_result,
            "new_score": new_score
        }

        return jsonify({"success": True, "data": response_data})

    except Exception as e:
        import traceback
        logger.error(f"Critical error in /api/v1/fix-issue-v2: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "An unexpected server error occurred."}), 500

# app.py में, /api/v1/generate-docx राउट को ढूंढो (लाइन 510-610 के आसपास)

@app.route('/api/v1/generate-docx', methods=['POST'])
def generate_docx_from_json():
    try:
        # 1. JSON डेटा को रिक्वेस्ट से निकालें
        # नोट: आपका प्रॉक्सी 'payload' को एक फॉर्म फील्ड में भेज रहा है
        if 'payload' not in request.form:
            return jsonify({"success": False, "error": "Missing payload"}), 400

        data = json.loads(request.form.get('payload'))
        
        doc = Document()

        # --- स्टाइलिंग (आप इसे और बेहतर बना सकते हैं) ---
        # फॉन्ट आदि यहाँ सेट करें
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # 2. नाम और टाइटल जोड़ें
        if data.get('name'):
            doc.add_heading(data['name'], level=1)
        if data.get('title'):
            p = doc.add_paragraph(data['title'])
            p.paragraph_format.space_after = Pt(18) # स्पेस डालें

        # 3. सभी सेक्शन्स को लूप करें और कंटेंट डालें
        for section in data.get('sections', []):
            if section.get('title'):
                doc.add_heading(section['title'].upper(), level=2)

            for content_item in section.get('content', []):
                content_type = content_item.get('type')

                if content_type == 'paragraph':
                    doc.add_paragraph(content_item.get('text', ''))

                elif content_type == 'list':
                    for item in content_item.get('items', []):
                        item_type = item.get('type')
                        
                        if item_type == 'experience':
                            # अनुभव वाले आइटम्स के लिए (हेडिंग, सब-हेडिंग, बुलेट पॉइंट्स)
                            p = doc.add_paragraph()
                            p.add_run(item.get('heading', '')).bold = True
                            if item.get('subheading'):
                                p.add_run(f"\n{item.get('subheading')}")
                            
                            for detail in item.get('details', []):
                                doc.add_paragraph(detail, style='List Bullet')

                        elif item_type == 'bullet':
                            # सरल बुलेट पॉइंट्स के लिए
                            doc.add_paragraph(item.get('text', ''), style='List Bullet')
            
            # हर सेक्शन के बाद थोड़ा स्पेस
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 4. फ़ाइल को मेमोरी में सेव करें और भेजें
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='AI_Generated_Resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": "Failed to generate DOCX file from JSON"}), 500 

# Function to download DOCX from backend (PDF will be client-side now, so no PDF logic here)
@app.route('/api/v1/markdown-to-docx', methods=['POST'])
def markdown_to_docx_handler():
    try:
        if 'payload' not in request.form:
            return jsonify({"success": False, "error": "Missing payload"}), 400

        data = json.loads(request.form.get('payload'))
        markdown_content = data.get("markdown_content")
        
        if not markdown_content:
            return jsonify({"success": False, "error": "Markdown content is empty"}), 400

        doc = Document()
        
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Handle Headings (H1, H2, H3)
            if line.startswith('###'):
                doc.add_heading(line.lstrip('# ').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.lstrip('# ').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.lstrip('# ').strip(), level=1)
            # Handle Bullet Points
            elif line.startswith('*'):
                p = doc.add_paragraph(line.lstrip('* ').strip(), style='List Bullet')
            # Handle everything else as a normal paragraph
            else:
                doc.add_paragraph(line)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='AI_Generated_Resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": "Failed to convert Markdown to DOCX"}), 500
    
@app.route('/api/v1/markdown-to-docx-fixer', methods=['POST'])
def markdown_to_docx_fixer_handler():
    try:
        if 'payload' not in request.form:
            return jsonify({"success": False, "error": "Missing payload"}), 400

        data = json.loads(request.form.get('payload'))
        markdown_content = data.get("markdown_content")
        
        if not markdown_content:
            return jsonify({"success": False, "error": "Markdown content is empty"}), 400

        doc = Document()
        
        for line in markdown_content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Handle Headings
            if line.startswith('###'):
                doc.add_heading(line.lstrip('# ').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.lstrip('# ').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.lstrip('# ').strip(), level=1)
            # Handle Bullet Points
            elif line.startswith('*'):
                doc.add_paragraph(line.lstrip('* ').strip(), style='List Bullet')
            # Handle everything else
            else:
                doc.add_paragraph(line)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='ResumeFixerPro_Resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": "Failed to convert Markdown for Fixer"}), 500
    
if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5000))
    logger.info(f"Starting Flask app on port {port}")
    app.run(host="0.0.0.0", port=port)

logger.info("Flask app initialization complete.")
